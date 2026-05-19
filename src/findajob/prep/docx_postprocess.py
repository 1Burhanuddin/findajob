"""Cosmetic post-processing for prep output documents.

- `_add_cover_letter_spacing(docx_path)` — applies 12pt paragraph spacing
  to body paragraphs in a cover letter `.docx`. The `BodyText` style in
  `config/reference.docx` has `space_after=20 twips` (~1pt), which renders
  too tight; direct paragraph formatting here overrides the style for the
  cover letter only.
- `_linkify_contact_info(md)` — converts bare email + LinkedIn URLs to
  Markdown hyperlinks before pandoc conversion so the docx has
  clickable links.

Imports are at module top so a missing `python-docx` install fails loudly
at orchestrator import time, not silently per-call. Cosmetic-only runtime
failures (corrupt docx, python-docx API drift) are caught and logged
without blocking prep.
"""

import re

from docx import Document
from docx.opc.exceptions import OpcError
from docx.shared import Pt

from findajob.audit import log_event


def _add_cover_letter_spacing(docx_path: str) -> None:
    """Post-process cover letter .docx for readable paragraph spacing.

    Heading 1 (paragraph 0) is left untouched — the reference.docx theme
    renders it correctly. Adjustments:
    1. Add 12pt space-before to paragraph [2] (the date) to separate it
       from the contact-info line.
    2. Apply 12pt space-after to every body paragraph from [2] onward,
       overriding the BodyText style's tight 1pt default.

    Bookmark anchors are no longer stripped here — `config/strip-bookmarks.lua`
    handles them at pandoc time, upstream of this post-process.
    """
    try:
        doc = Document(docx_path)
        if not doc.paragraphs:
            return

        if len(doc.paragraphs) > 2:
            doc.paragraphs[2].paragraph_format.space_before = Pt(12)

        for para in doc.paragraphs[2:]:
            if para.text.strip():
                para.paragraph_format.space_after = Pt(12)

        doc.save(docx_path)
    except (KeyError, AttributeError, ValueError, OSError, OpcError) as exc:
        log_event(
            "cover_letter_spacing_failed",
            docx_path=docx_path,
            error=f"{type(exc).__name__}: {exc}",
        )


def _linkify_contact_info(md: str) -> str:
    """Ensure bare email addresses and LinkedIn URLs are Markdown hyperlinks.

    Runs on resume markdown before pandoc conversion so the .docx has clickable links.
    Skips anything already inside []() link syntax.
    """
    # Email: bare user@domain.tld → [user@domain.tld](mailto:user@domain.tld)
    md = re.sub(
        r"(?<!\[)(?<!\(mailto:)\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})\b(?!\])",
        r"[\1](mailto:\1)",
        md,
    )
    # LinkedIn: bare linkedin.com/in/handle → [linkedin.com/in/handle](https://linkedin.com/in/handle)
    md = re.sub(
        r"(?<!\[)(?<!\(https://)(linkedin\.com/in/[A-Za-z0-9_-]+)(?!\])",
        r"[\1](https://\1)",
        md,
    )
    return md
