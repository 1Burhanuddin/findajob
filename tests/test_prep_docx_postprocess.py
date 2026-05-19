"""Tests for findajob.prep.docx_postprocess — cover letter spacing fix (#739).

Background: `_add_cover_letter_spacing` is the regression site. Before #739,
its `from docx import Document` was inside a bare `try/except Exception: pass`
block at the function top — and `python-docx` was missing from
`pyproject.toml` dependencies. ImportError was silently swallowed on every
invocation, so the post-processor had been a no-op the entire time. Body
paragraphs inherited `BodyText` style's `space_after=20 twips` (~1pt),
rendering "tight" instead of with readable blank-line gaps.

The fix: add `python-docx` to deps, hoist imports to module top (so missing
dep fails loudly at import time, not silently per-call), and narrow the
in-function except to specific runtime errors.

Test discipline (advisor's call in #739): the test must fail in a clean
venv that doesn't have `python-docx` installed. The module-top import is
what enforces that — collecting this test file requires importing
`findajob.prep.docx_postprocess`, which requires `docx`.
"""

from __future__ import annotations

from docx import Document
from docx.shared import Pt

from findajob.prep.docx_postprocess import _add_cover_letter_spacing


def _build_cover_letter_docx(path: str) -> None:
    """Build a minimal cover-letter-shaped .docx fixture.

    Mirrors what pandoc emits for the canonical cover letter structure:
    paragraph [0] is the Heading 1, [1] is the contact info (FirstParagraph
    style), and [2:] are BodyText paragraphs (date, hiring team, Re:, body,
    sign-off).
    """
    doc = Document()
    doc.add_heading("CoreWeave | Senior Program Manager", level=1)
    doc.add_paragraph("Alex Doe · Anytown, USA · 555-555-5555")
    doc.add_paragraph("May 19, 2026")
    doc.add_paragraph("CoreWeave Hiring Team")
    doc.add_paragraph("Re: Senior Program Manager")
    doc.add_paragraph("This is the first body paragraph.")
    doc.add_paragraph("This is the second body paragraph.")
    doc.add_paragraph("Alex Doe")
    doc.save(path)


def test_spacing_applied_to_body_paragraphs(tmp_path):
    """Body paragraphs [2:] must end with space_after >= Pt(12) after post-process."""
    docx_path = str(tmp_path / "cover.docx")
    _build_cover_letter_docx(docx_path)

    _add_cover_letter_spacing(docx_path)

    doc = Document(docx_path)
    for i, para in enumerate(doc.paragraphs[2:], start=2):
        if not para.text.strip():
            continue
        space_after = para.paragraph_format.space_after
        assert space_after is not None, f"paragraph [{i}] has no space_after override"
        assert space_after >= Pt(12), (
            f"paragraph [{i}] space_after={space_after} < Pt(12); this is the regression that #739 fixed"
        )


def test_heading_and_contact_paragraphs_left_untouched(tmp_path):
    """The heading and contact-info paragraphs must not have space_after set."""
    docx_path = str(tmp_path / "cover.docx")
    _build_cover_letter_docx(docx_path)

    _add_cover_letter_spacing(docx_path)

    doc = Document(docx_path)
    assert doc.paragraphs[0].paragraph_format.space_after is None, (
        "heading [0] space_after should be inherited from style, not overridden"
    )
    assert doc.paragraphs[1].paragraph_format.space_after is None, (
        "contact-info [1] space_after should be inherited from style, not overridden"
    )


def test_date_paragraph_gets_space_before(tmp_path):
    """Paragraph [2] (the date) gets a space_before override to separate from contact info."""
    docx_path = str(tmp_path / "cover.docx")
    _build_cover_letter_docx(docx_path)

    _add_cover_letter_spacing(docx_path)

    doc = Document(docx_path)
    space_before = doc.paragraphs[2].paragraph_format.space_before
    assert space_before is not None
    assert space_before >= Pt(12)


def test_empty_paragraphs_not_modified(tmp_path):
    """Paragraphs with no visible text must be skipped (post-process is conditional on .text.strip())."""
    doc = Document()
    doc.add_heading("Test | Role", level=1)
    doc.add_paragraph("Contact line")
    doc.add_paragraph("")  # empty paragraph at index [2]
    doc.add_paragraph("Body text")
    docx_path = str(tmp_path / "cover.docx")
    doc.save(docx_path)

    _add_cover_letter_spacing(docx_path)

    doc2 = Document(docx_path)
    assert doc2.paragraphs[2].paragraph_format.space_after is None
    assert doc2.paragraphs[3].paragraph_format.space_after == Pt(12)


def test_missing_file_logs_event_does_not_raise(tmp_path):
    """Cosmetic post-process must never block prep — missing file is caught + logged."""
    nonexistent = str(tmp_path / "does-not-exist.docx")

    # Must not raise; an OSError from python-docx open is the expected fail path.
    _add_cover_letter_spacing(nonexistent)
