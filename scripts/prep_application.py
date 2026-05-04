#!/usr/bin/env python3
# scripts/prep_application.py
# Args: company, title, url, job_id
"""Generate draft application materials for a flagged job.

Launched as a detached subprocess from POST /board/jobs/{fp}/prep (see
findajob.web.routes.board_actions).
"""

import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
from datetime import UTC, datetime

from findajob.actions import reset_prep_to_scored
from findajob.paths import AICHAT, BASE, PANDOC
from findajob.utils import (
    JD_MAX_CHARS,
    build_prep_filenames,
    load_env,
    load_voice_samples,
    log_event,
    quarantine_stale_prep_folders,
    read_file_prefix,
    write_audit,
)

DB_PATH = f"{BASE}/data/pipeline.db"
PROFILE_PATH = f"{BASE}/candidate_context/profile.md"
MASTER_RESUME_PATH = f"{BASE}/candidate_context/master_resume.md"

load_env()


def aichat(role, prompt, model_override=None, timeout=300):
    """Call aichat-ng and return stdout. No RAG — all context injected directly."""
    cmd = [AICHAT, "--role", role]
    if model_override:
        cmd += ["-m", model_override]
    cmd += ["-S", prompt]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    output = result.stdout.strip()
    if result.returncode != 0 or not output:
        log_event("aichat_failure", role=role, returncode=result.returncode, stderr=result.stderr.strip()[:500])
    # Strip <think>...</think> blocks that leak from :thinking models
    output = re.sub(r"<think>.*?</think>", "", output, flags=re.DOTALL).strip()
    return output


def abbrev_title(title, max_words=3):
    """Return a folder-safe abbreviated title: first N significant words joined with underscores."""
    title = re.sub(r"\s*\(.*?\)", "", title)  # strip parentheticals
    title = re.sub(r"[^\w\s-]", "", title)  # remove punctuation
    words = [w for w in title.split() if w][:max_words]
    return "_".join(words) if words else "Job"


def _add_cover_letter_spacing(docx_path):
    """Post-process cover letter .docx for clean formatting.

    Heading 1 is left untouched — the reference.docx theme renders it correctly
    (teal color, heading font) in Google Docs. Adjustments:
    1. Remove pandoc bookmark anchors (render as blue bracket in Google Docs)
    2. Space before the date line to separate from contact info
    3. 12pt space-after from date onward for readable paragraph gaps
    """
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document(docx_path)
        if not doc.paragraphs:
            return

        # 1. Strip bookmark anchors that pandoc adds to headings
        body = doc.element.body
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for tag in ("bookmarkStart", "bookmarkEnd"):
            for bm in body.findall(f".//{{{ns}}}{tag}"):
                bm.getparent().remove(bm)

        # 2. Add space before the date line (paragraph [2]) to separate from contact
        if len(doc.paragraphs) > 2:
            doc.paragraphs[2].paragraph_format.space_before = Pt(12)

        # 3. Add 12pt space-after from date onward
        for para in doc.paragraphs[2:]:
            if para.text.strip():
                para.paragraph_format.space_after = Pt(12)

        doc.save(docx_path)
    except Exception:
        pass  # post-processing is cosmetic — never block prep


def _linkify_contact_info(md):
    """Ensure bare email addresses and LinkedIn URLs are Markdown hyperlinks.

    Runs on resume markdown before pandoc conversion so the .docx has clickable links.
    Skips anything already inside []() link syntax.
    """
    # Email: bare user@domain.tld → [user@domain.tld](mailto:user@domain.tld)
    md = re.sub(
        r"(?<!\[)(?<!\(mailto:)\b([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})\b(?!\])",
        r"[\1](mailto:\1)",
        md,
    )
    # LinkedIn: bare linkedin.com/in/handle → [linkedin.com/in/handle](https://linkedin.com/in/handle)
    md = re.sub(
        r"(?<!\[)(?<!\(https://)(linkedin\.com/in/[A-Za-z0-9_-]+)(?!\])",
        r"[\1](https://\1)",
        md,
    )
    return md


def notify(message):
    topic = None
    try:
        with open(f"{BASE}/config/ntfy_topic.txt") as f:
            topic = f.read().strip()
    except FileNotFoundError:
        pass
    if not topic:
        # Fall back to data/.env NTFY_TOPIC
        try:
            with open(f"{BASE}/data/.env") as f:
                for line in f:
                    line = line.strip()
                    if line.startswith("NTFY_TOPIC") and "=" in line:
                        topic = line.split("=", 1)[1].strip().strip("'\"")
                        break
        except Exception:
            pass
    if not topic:
        return
    try:
        subprocess.run(["curl", "-s", "-d", message, f"https://ntfy.sh/{topic}"], capture_output=True, timeout=10)
    except Exception:
        pass


def main():
    company, title, url, job_id = sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4]

    # Guard: skip if prep already completed for this job
    conn_check = sqlite3.connect(DB_PATH, timeout=30)
    conn_check.row_factory = sqlite3.Row
    existing = conn_check.execute("SELECT prep_folder_path, stage FROM jobs WHERE id=?", (job_id,)).fetchone()
    conn_check.close()
    if existing and existing["prep_folder_path"] and existing["stage"] == "materials_drafted":
        log_event("prep_skipped_duplicate", company=company, title=title, job_id=job_id)
        print(f"PREP_SKIPPED: materials already drafted for {job_id}")
        return

    # Sanitize company for filesystem safety (title already goes through abbrev_title)
    safe_company = re.sub(r"[^\w\s\-&.,]", "_", company).strip()
    date = datetime.now().strftime("%Y-%m-%d")
    time_str = datetime.now().strftime("%H%M%S")
    companies_dir = f"{BASE}/companies"
    folder_prefix = f"{safe_company}_{abbrev_title(title)}_"
    outdir = f"{companies_dir}/{folder_prefix}{date}_{time_str}"

    # Quarantine any prior prep folders for this {company, title} that aren't
    # tracked by the DB — Regenerate and prep races otherwise leave orphans (#174).
    cleanup_conn = sqlite3.connect(DB_PATH, timeout=30)
    try:
        quarantine_stale_prep_folders(cleanup_conn, companies_dir, folder_prefix, os.path.basename(outdir))
    finally:
        cleanup_conn.close()

    os.makedirs(outdir, exist_ok=True)

    # Build per-file output paths using the candidate's file prefix (from profile.md).
    # Pattern: {Prefix} Resume - {Company} - {Title} - {YYYYMMDD-HHMMSS}.{ext}
    # See scripts/utils.py:build_prep_filenames for the full pattern.
    file_prefix = read_file_prefix()
    timestamp_fn = f"{date.replace('-', '')}-{time_str}"
    fn = build_prep_filenames(company, title, timestamp_fn, file_prefix)
    out = {k: os.path.join(outdir, v) for k, v in fn.items()}

    log_event("prep_started", company=company, title=title, job_id=job_id, file_prefix=file_prefix)

    # ── Step 1: Load JD from DB (already fetched during triage) ──
    # Do NOT re-curl — LinkedIn and many other URLs require auth and will return garbage.
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    row = conn.execute(
        "SELECT raw_jd_text, stage, synthetic, speculative_briefing_folder FROM jobs WHERE id=?",
        (job_id,),
    ).fetchone()
    jd_text = (row["raw_jd_text"] or "").strip() if row else ""
    is_synthetic = bool(row["synthetic"]) if row and "synthetic" in row.keys() else False
    mode_marker = "<<SPECULATIVE_MODE>>\n\n" if is_synthetic else ""
    speculative_briefing_folder = (
        row["speculative_briefing_folder"] if row and "speculative_briefing_folder" in row.keys() else None
    )

    if not jd_text or len(jd_text) < 50:
        # Fallback: try curling for Greenhouse/Lever/public URLs only
        try:
            raw = subprocess.run(["curl", "-sL", "--max-time", "15", url], capture_output=True, text=True).stdout
            jd_text = subprocess.run(
                [PANDOC, "-f", "html", "-t", "plain"], input=raw, capture_output=True, text=True
            ).stdout[:JD_MAX_CHARS]
        except Exception:
            jd_text = "[ERROR: Could not fetch JD]"

    with open(out["jd_txt"], "w") as f:
        f.write(jd_text)

    # ── Load profile and master resume — injected directly, never via RAG ──
    missing_files = []
    try:
        with open(PROFILE_PATH) as f:
            profile_text = f.read()
    except FileNotFoundError:
        missing_files.append(f"profile.md ({PROFILE_PATH})")
        profile_text = ""

    try:
        with open(MASTER_RESUME_PATH) as f:
            master_text = f.read()
    except FileNotFoundError:
        missing_files.append(f"master_resume.md ({MASTER_RESUME_PATH})")
        master_text = ""

    if missing_files:
        log_event("prep_missing_candidate_files", job_id=job_id, company=company, missing="; ".join(missing_files))
        shutil.rmtree(outdir, ignore_errors=True)
        reset_prep_to_scored(conn, job_id, reason="missing_candidate_files")
        notify(f"PREP ABORTED (missing candidate files): {company} — {title}\n{'; '.join(missing_files)}")
        return

    # ── Step 2: Company briefing FIRST — gives all downstream steps rich context ──
    # For synthetic rows (#131 speculative), the deep-research briefing was
    # already generated at submission time and approved by the operator on the
    # review page. Reuse it instead of regenerating via briefing_writer (#320 —
    # spec drift fix). Falls back to the regular briefing_writer flow if the
    # column is unset, the folder is missing, or briefing.md is empty/absent.
    rec_re = re.compile(r"^##[^\n]*Overall Recommendation\s*:", re.MULTILINE)
    briefing = ""
    if is_synthetic and speculative_briefing_folder:
        spec_briefing_path = os.path.join(BASE, "companies", speculative_briefing_folder, "briefing.md")
        try:
            with open(spec_briefing_path) as f:
                briefing = f.read().strip()
            if briefing:
                log_event(
                    "speculative_briefing_reused",
                    job_id=job_id,
                    company=company,
                    folder=speculative_briefing_folder,
                    chars=len(briefing),
                )
        except FileNotFoundError:
            log_event(
                "speculative_briefing_missing",
                job_id=job_id,
                company=company,
                folder=speculative_briefing_folder,
                expected_path=spec_briefing_path,
            )
            briefing = ""

    if not briefing:
        # Real-row flow OR synthetic-fallback when the speculative briefing is missing.
        brief_prompt = f"Research {company} thoroughly.\nJob title: {title}\nJD:\n{jd_text}"
        raw_briefing = aichat("company_researcher", brief_prompt)

        # Pass raw research through briefing_writer with candidate context for stories
        formatted_brief_prompt = (
            f"Format the following company research into a structured briefing 1-pager "
            f"for {company}. Job: {title}.\n\n"
            f"RAW RESEARCH:\n{raw_briefing}\n\n"
            f"CANDIDATE PROFILE:\n{profile_text}\n\n"
            f"MASTER RESUME:\n{master_text}\n\n"
            f"JD:\n{jd_text}"
        )
        briefing = aichat("briefing_writer", formatted_brief_prompt)

        # ── Validate: briefing must end with an Overall Recommendation section ──
        # The role prompt requires this verdict heading; model sometimes drops it.
        # Retry once; if still missing, let downstream validator fail prep cleanly.
        if not briefing or not rec_re.search(briefing):
            log_event("briefing_missing_recommendation", job_id=job_id, company=company, retry=1)
            briefing = aichat("briefing_writer", formatted_brief_prompt)

    # Fit analysis: multi-dimensional assessment appended to briefing
    fit_prompt = (
        f"Analyze the fit between this candidate and this role.\n\n"
        f"CANDIDATE PROFILE:\n{profile_text}\n\n"
        f"MASTER RESUME:\n{master_text}\n\n"
        f"Company: {company}\nTitle: {title}\n\n"
        f"JD:\n{jd_text}\n\n"
        f"COMPANY BRIEFING:\n{briefing}"
    )
    fit_analysis = aichat("fit_analyst", fit_prompt)

    # Combine briefing and fit analysis into one document.
    # The briefing ends with an Overall Recommendation verdict; fit analysis
    # contains the Matrix/Probability/Strengths/Gaps detail that should sit
    # BEFORE the verdict so the doc reads detail → synthesis → recommendation.
    fit_score_avg = None
    prob_score_avg = None
    full_briefing = briefing
    if fit_analysis:
        rec_match = rec_re.search(briefing)
        if rec_match:
            briefing_pre = briefing[: rec_match.start()].rstrip()
            briefing_rec = briefing[rec_match.start() :]
            full_briefing = f"{briefing_pre}\n\n---\n\n# Fit Analysis\n\n{fit_analysis}\n\n---\n\n{briefing_rec}"
        else:
            full_briefing = f"{briefing}\n\n---\n\n# Fit Analysis\n\n{fit_analysis}"
        # Parse scores from fit analysis for DB storage
        # All scores are 0-100%. Fit Matrix section has 6 dimensions, Probability has 3.
        try:
            # Split on Probability Assessment heading to separate the two sections
            parts = re.split(r"##\s*🎯\s*Probability Assessment", fit_analysis, maxsplit=1)
            fit_section = parts[0] if parts else fit_analysis
            prob_section = parts[1] if len(parts) > 1 else ""
            fit_scores = [int(m.group(1)) for m in re.finditer(r":\s*(\d{1,3})%", fit_section)]
            prob_scores = [int(m.group(1)) for m in re.finditer(r":\s*(\d{1,3})%", prob_section)]
            if fit_scores:
                fit_score_avg = round(sum(fit_scores) / len(fit_scores), 1)
            if prob_scores:
                prob_score_avg = round(sum(prob_scores) / len(prob_scores), 1)
            log_event(
                "fit_analysis",
                company=company,
                title=title,
                fit_score=fit_score_avg,
                probability_score=prob_score_avg,
                fit_scores=fit_scores,
                prob_scores=prob_scores,
            )
        except Exception:
            pass

    with open(out["briefing_md"], "w") as f:
        f.write(full_briefing)
    subprocess.run(
        [
            PANDOC,
            "-f",
            "markdown-yaml_metadata_block",
            out["briefing_md"],
            "--lua-filter",
            f"{BASE}/config/strip-bookmarks.lua",
            "--reference-doc",
            f"{BASE}/config/reference.docx",
            "-o",
            out["briefing_docx"],
        ],
        check=False,
    )

    # ── Step 3: Resume — briefing + fit analysis context now available ──
    briefing_context = full_briefing if full_briefing else ""
    resume_prompt = (
        f"MASTER RESUME:\n{master_text}\n\n"
        f"CANDIDATE PROFILE:\n{profile_text}\n\n"
        f"Company: {company}\nTitle: {title}\n\n"
        f"JD:\n{jd_text}\n\n"
        f"COMPANY BRIEFING AND FIT ANALYSIS:\n{briefing_context}"
    )
    resume_md = aichat("resume_tailor", resume_prompt)
    # Strip [VERIFY: ...] lines that appear before the first # header
    rlines = resume_md.split("\n")
    first_hdr = next((i for i, line in enumerate(rlines) if line.startswith("#")), 0)
    rlines = [line for i, line in enumerate(rlines) if not (i < first_hdr and line.startswith("[VERIFY:"))]
    resume_md = "\n".join(rlines).strip()
    resume_md = _linkify_contact_info(resume_md)
    with open(out["resume_md"], "w") as f:
        f.write(resume_md)

    # Quality check — log violation counts for trend tracking
    try:
        qc = subprocess.run(
            [sys.executable, f"{BASE}/scripts/diag/validate_resume.py", "--json", out["resume_md"]],
            capture_output=True,
            text=True,
            timeout=15,
        )
        if qc.stdout:
            viols_by_file = json.loads(qc.stdout)
            viols = list(viols_by_file.values())[0] if viols_by_file else []
            log_event(
                "resume_quality_check",
                company=company,
                title=title,
                violations=len(viols),
                high=sum(1 for v in viols if v["severity"] == "HIGH"),
                med=sum(1 for v in viols if v["severity"] == "MED"),
            )
    except Exception:
        pass  # quality check is informational only — never block prep

    subprocess.run(
        [
            PANDOC,
            out["resume_md"],
            "--lua-filter",
            f"{BASE}/config/strip-bookmarks.lua",
            "--reference-doc",
            f"{BASE}/config/reference.docx",
            "-o",
            out["resume_docx"],
        ],
        check=False,
    )

    # Generate change log
    changes_prompt = f"ORIGINAL MASTER RESUME:\n{master_text}\n\nTAILORED RESUME:\n{resume_md}\n\nTARGET JD:\n{jd_text}"
    changes_md = aichat("resume_change_reviewer", changes_prompt)
    with open(out["changes_md"], "w") as f:
        f.write(changes_md)

    # ── Step 4: Cover letter — briefing + fit analysis for company signals ──
    today_str = datetime.now().strftime("%B %d, %Y")
    voice_samples = load_voice_samples()
    log_event("voice_samples_loaded", caller="cover_letter_writer", chars=len(voice_samples))
    voice_section = f"VOICE SAMPLES:\n{voice_samples}\n\n" if voice_samples else ""
    cover_prompt = (
        f"{mode_marker}"
        f"CANDIDATE PROFILE:\n{profile_text}\n\n"
        f"MASTER RESUME:\n{master_text}\n\n"
        f"{voice_section}"
        f"Company: {company}\nTitle: {title}\nDate: {today_str}\n\n"
        f"JD:\n{jd_text}\n\n"
        f"COMPANY BRIEFING AND FIT ANALYSIS:\n{briefing_context}"
    )
    cover_md_text = aichat("cover_letter_writer", cover_prompt)
    # Strip horizontal rules — the LLM inserts "---" between header and body,
    # but it renders as an ugly line in the docx. Paragraph spacing handles separation.
    cover_md_text = re.sub(r"\n---\n", "\n\n", cover_md_text)
    with open(out["cover_md"], "w") as f:
        f.write(cover_md_text)
    subprocess.run(
        [
            PANDOC,
            out["cover_md"],
            "--lua-filter",
            f"{BASE}/config/strip-bookmarks.lua",
            "--reference-doc",
            f"{BASE}/config/reference.docx",
            "-o",
            out["cover_docx"],
        ],
        check=False,
    )
    _add_cover_letter_spacing(out["cover_docx"])

    # ── Step 4.5: Recruiter critique — skeptical outside read of resume + cover ──
    # Sees only what an actual recruiter sees: company, title, JD, resume, cover.
    # No profile / briefing / fit analysis — the point is to simulate a reader who
    # has NOT done background research on the candidate.
    critique_prompt = (
        f"Company: {company}\nTitle: {title}\n\n"
        f"JD:\n{jd_text}\n\n"
        f"TAILORED RESUME:\n{resume_md}\n\n"
        f"COVER LETTER:\n{cover_md_text}"
    )
    critique_md = aichat("recruiter_critic", critique_prompt)
    if critique_md:
        with open(out["critique_md"], "w") as f:
            f.write(critique_md)

    # ── Step 5: Network outreach ──
    # Pass the file_prefix and timestamp so outreach files follow the same naming convention.
    subprocess.run(
        [
            sys.executable,
            f"{BASE}/scripts/find_contacts.py",
            company,
            jd_text,
            outdir,
            file_prefix,
            timestamp_fn,
            "1" if is_synthetic else "0",
        ],
        check=False,
    )

    # ── Step 6: Review checklist ──
    with open(out["checklist_md"], "w") as f:
        f.write(f"""# Review Checklist — {company} / {title}
Generated: {date}

## Before sending, complete these steps:
- [ ] Open `{fn["changes_md"]}` — review every flagged reorder/keyword add
- [ ] Open `{fn["resume_docx"]}` — fill any [MISSING: ...] placeholders
- [ ] Open `{fn["cover_docx"]}` — fill any [MISSING: ...] placeholders (expect 1-2 max)
- [ ] Read cover letter aloud — does it sound like you?
- [ ] Verify every factual claim in the cover letter (metrics, company names, titles)
- [ ] Check `{fn["briefing_docx"]}` — any red flags or new intel to weave in?
- [ ] Review outreach drafts if you plan to reach out before applying

## Files in this folder:
- `{fn["resume_docx"]}`    ← start here
- `{fn["changes_md"]}`    ← what the AI changed and why
- `{fn["cover_docx"]}`       ← fill placeholders before sending
- `{fn["briefing_docx"]}`
- `{fn["jd_txt"]}`    ← original JD for reference
- `{file_prefix} Outreach to *.txt`    ← network outreach drafts
""")

    # ── Step 7: Validate output before marking complete ──
    # Guard against silent LLM failures (e.g. missing max_tokens) that produce
    # empty files.  On failure: delete the damaged folder, reset to scored so
    # the job can be re-prepped, and abort.
    MIN_BYTES = 500
    validation_failures = []
    for label, path in [("resume", out["resume_md"]), ("cover_letter", out["cover_md"])]:
        try:
            sz = os.path.getsize(path)
        except OSError:
            sz = 0
        if sz < MIN_BYTES:
            validation_failures.append(f"{label}: {sz}B (min {MIN_BYTES})")
    # Briefing must end with an Overall Recommendation verdict — model drift
    # sometimes drops it despite role prompt enforcement.
    try:
        with open(out["briefing_md"]) as f:
            briefing_text = f.read()
    except OSError:
        briefing_text = ""
    if not rec_re.search(briefing_text):
        validation_failures.append("briefing: missing Overall Recommendation")
    if validation_failures:
        log_event(
            "prep_validation_failed",
            company=company,
            title=title,
            failures="; ".join(validation_failures),
        )
        shutil.rmtree(outdir, ignore_errors=True)
        reset_prep_to_scored(conn, job_id, reason="validation_failed")
        notify(f"PREP FAILED (empty files): {company} — {title}\n{'; '.join(validation_failures)}")
        return

    # ── Step 8: Update SQLite (stage + scores) ──
    now = datetime.now(UTC).isoformat()
    old_stage = row["stage"] if row else "unknown"

    conn.execute(
        """
        UPDATE jobs SET stage='materials_drafted', stage_updated=?, prep_folder_path=?,
               fit_score=?, probability_score=?, updated_at=?
        WHERE id=?
    """,
        (now, outdir, fit_score_avg, prob_score_avg, now, job_id),
    )
    conn.commit()
    write_audit(conn, job_id, "stage", old_stage, "materials_drafted")

    log_event("prep_complete", company=company, title=title, folder=outdir)
    notify(f"Drafts ready: {company} — {title}\n{outdir}")

    conn.close()

    print(f"PREP_COMPLETE:{outdir}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        # Recover from any unhandled error: log failure and reset stage so
        # the job can be retried on the next poll cycle.
        job_id = sys.argv[4] if len(sys.argv) > 4 else "unknown"
        company = sys.argv[1] if len(sys.argv) > 1 else "unknown"
        title = sys.argv[2] if len(sys.argv) > 2 else "unknown"
        log_event(
            "prep_failed",
            job_id=job_id,
            company=company,
            title=title,
            error=f"{type(exc).__name__}: {exc}",
        )
        try:
            conn = sqlite3.connect(DB_PATH, timeout=30)
            reset_prep_to_scored(conn, job_id, reason=f"exception:{type(exc).__name__}")
            conn.close()
        except Exception:
            pass  # DB recovery is best-effort
        raise
